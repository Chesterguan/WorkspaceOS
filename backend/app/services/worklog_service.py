"""Work Log service — gather project data, generate reports, export DOCX."""
import io
import logging
import re
from datetime import date
from typing import Any, Dict, List, Optional

from sqlalchemy import text
from sqlalchemy.ext.asyncio import AsyncSession

from app.services.ai_client import get_cloud_client
from app.services.egress_recorder import EgressRecorder

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Data gathering
# ---------------------------------------------------------------------------

async def gather_period_data(
    project_ids: List,
    start_date: date,
    end_date: date,
    db: AsyncSession,
) -> Dict[str, Any]:
    """Query commits, papers, drafts, and syncs for the given projects + date range."""
    pid_list = [str(p) for p in project_ids]

    # Project names + pinned focus_notes in one query.
    rows = await db.execute(
        text("SELECT id, name, focus_notes FROM projects WHERE id = ANY(:ids)"),
        {"ids": pid_list},
    )
    project_names: Dict[str, str] = {}
    focus_notes: Dict[str, Optional[str]] = {}
    for r in rows.fetchall():
        project_names[str(r[0])] = r[1]
        focus_notes[str(r[0])] = r[2]

    # Latest wiki_summary per project — the evolving "what is this project"
    # narrative already maintained by memory_service.upsert_wiki_summary.
    # One entry per project by construction, but DISTINCT ON guards against
    # duplicates if that invariant ever slips.
    rows = await db.execute(
        text("""
            SELECT DISTINCT ON (project_id) project_id, content
            FROM memory_entries
            WHERE project_id = ANY(:ids)
              AND entry_type = 'wiki_summary'
            ORDER BY project_id, created_at DESC
        """),
        {"ids": pid_list},
    )
    wiki_by_project: Dict[str, str] = {str(r[0]): r[1] for r in rows.fetchall()}

    # Total commits per project
    rows = await db.execute(
        text("""
            SELECT project_id, COUNT(*) AS cnt
            FROM github_commits
            WHERE project_id = ANY(:ids)
              AND committed_at >= :start AND committed_at <= :end
            GROUP BY project_id
        """),
        {"ids": pid_list, "start": start_date, "end": end_date},
    )
    commits_by_project = {str(r[0]): r[1] for r in rows.fetchall()}

    # Weekly commit breakdown (for chart)
    rows = await db.execute(
        text("""
            SELECT date_trunc('week', committed_at)::date AS week_start,
                   COUNT(*) AS cnt
            FROM github_commits
            WHERE project_id = ANY(:ids)
              AND committed_at >= :start AND committed_at <= :end
            GROUP BY week_start
            ORDER BY week_start
        """),
        {"ids": pid_list, "start": start_date, "end": end_date},
    )
    weekly_commits = [
        {"week": str(r[0]), "count": r[1]} for r in rows.fetchall()
    ]

    # Papers (blog_posts with tags containing 'paper')
    rows = await db.execute(
        text("""
            SELECT title, created_at
            FROM blog_posts
            WHERE project_id = ANY(:ids)
              AND tags @> ARRAY['paper']
              AND created_at >= :start AND created_at <= :end
            ORDER BY created_at
        """),
        {"ids": pid_list, "start": start_date, "end": end_date},
    )
    papers = [{"title": r[0], "created_at": str(r[1])} for r in rows.fetchall()]

    # Drafts
    rows = await db.execute(
        text("""
            SELECT status, COUNT(*) AS cnt
            FROM drafts
            WHERE project_id = ANY(:ids)
              AND created_at >= :start AND created_at <= :end
            GROUP BY status
        """),
        {"ids": pid_list, "start": start_date, "end": end_date},
    )
    drafts_by_status = {r[0]: r[1] for r in rows.fetchall()}

    # Sync runs
    rows = await db.execute(
        text("""
            SELECT project_id, status, commits_fetched, completed_at
            FROM sync_runs
            WHERE project_id = ANY(:ids)
              AND completed_at >= :start AND completed_at <= :end
            ORDER BY completed_at
        """),
        {"ids": pid_list, "start": start_date, "end": end_date},
    )
    syncs = [
        {
            "project": project_names.get(str(r[0]), str(r[0])),
            "status": r[1],
            "commits_fetched": r[2],
            "completed_at": str(r[3]),
        }
        for r in rows.fetchall()
    ]

    # Knowledge nodes created in this period (decisions/questions/blockers/etc).
    rows = await db.execute(
        text("""
            SELECT project_id, node_type, title, content, created_at
            FROM knowledge_nodes
            WHERE project_id = ANY(:ids)
              AND archived = false
              AND created_at >= :start AND created_at <= :end
            ORDER BY created_at
        """),
        {"ids": pid_list, "start": start_date, "end": end_date},
    )
    knowledge_by_project: Dict[str, Dict[str, list]] = {}
    for r in rows.fetchall():
        pid = str(r[0])
        bucket = knowledge_by_project.setdefault(pid, {})
        bucket.setdefault(r[1], []).append({
            "title": r[2], "content": r[3], "created_at": str(r[4]),
        })

    # project_context: per-project narrative baseline the LLM should treat as
    # authoritative (user-pinned focus + AI-maintained wiki). Empty when a
    # project has no pinned notes and no wiki yet — prompt section still
    # renders but stays terse.
    project_context: Dict[str, Dict[str, Optional[str]]] = {
        pid: {
            "name": project_names.get(pid, pid),
            "focus": focus_notes.get(pid),
            "wiki": wiki_by_project.get(pid),
        }
        for pid in pid_list
    }

    return {
        "project_names": project_names,
        "commits_by_project": commits_by_project,
        "weekly_commits": weekly_commits,
        "papers": papers,
        "drafts_by_status": drafts_by_status,
        "syncs": syncs,
        "project_context": project_context,
        "knowledge_by_project": knowledge_by_project,
        "period_start": str(start_date),
        "period_end": str(end_date),
    }


# ---------------------------------------------------------------------------
# Report generation
# ---------------------------------------------------------------------------

# Cap the wiki-per-project slice to keep the prompt under control when the
# report covers many projects. Tuned to leave room for focus + metrics in a
# typical 8k-token budget.
_WIKI_CHARS_PER_PROJECT = 1500


def _build_context_block(project_context: Dict[str, Dict[str, Optional[str]]]) -> str:
    """Render per-project focus + wiki as a markdown section for the prompt."""
    if not project_context:
        return ""
    parts: List[str] = ["## Project Context"]
    for ctx in project_context.values():
        parts.append(f"### {ctx.get('name') or 'Unnamed project'}")
        focus = (ctx.get("focus") or "").strip()
        parts.append(f"**Current focus (user-pinned):** {focus or '—'}")
        wiki = (ctx.get("wiki") or "").strip()
        if wiki:
            if len(wiki) > _WIKI_CHARS_PER_PROJECT:
                wiki = wiki[:_WIKI_CHARS_PER_PROJECT] + "\n… (truncated)"
            parts.append("**Project wiki summary:**")
            parts.append(wiki)
        else:
            parts.append("**Project wiki summary:** (no wiki generated yet)")
    return "\n".join(parts)


async def generate_report(
    period_type: str,
    period_data: Dict[str, Any],
    goals: Optional[List[Dict[str, str]]] = None,
    additional_instructions: Optional[str] = None,
) -> str:
    """Call cloud AI to generate the progress report markdown."""
    from app.services.domain_config import get_loader

    try:
        system_prompt = get_loader().get_worklog_template(period_type)
    except KeyError:
        system_prompt = get_loader().get_worklog_template("weekly")

    user_parts: List[str] = [
        f"Period: {period_data['period_start']} to {period_data['period_end']}",
        f"Projects: {', '.join(period_data['project_names'].values())}",
    ]

    # Project Context goes first so the LLM anchors on the narrative before
    # seeing the metrics. Rendered only when we actually have context to
    # show — prevents a lonely "## Project Context" header over nothing.
    context_block = _build_context_block(period_data.get("project_context") or {})
    if context_block:
        user_parts.append("")
        user_parts.append(context_block)

    # Render knowledge nodes — group by node_type globally, then by project.
    knowledge_by_project = period_data.get("knowledge_by_project") or {}
    if knowledge_by_project:
        kn_lines: List[str] = ["## Knowledge Captured This Period"]
        # Group by node_type across projects so the supervisor sees patterns.
        by_type: Dict[str, List[str]] = {}
        for pid, types_dict in knowledge_by_project.items():
            project_label = period_data["project_names"].get(pid, pid)
            for ntype, items in types_dict.items():
                bucket = by_type.setdefault(ntype, [])
                for it in items:
                    bucket.append(f"- ({project_label}) **{it['title']}** — {it['content'][:300]}")
        type_order = ["decision", "rejection", "question", "blocker",
                      "claim", "hypothesis", "insight"]
        for ntype in type_order:
            items = by_type.get(ntype)
            if items:
                kn_lines.append(f"### {ntype.title()}s")
                kn_lines.extend(items)
        if len(kn_lines) > 1:
            user_parts.append("\n".join(kn_lines))

    user_parts.extend([
        "",
        "## Period Metrics",
        f"Commits by project: {period_data['commits_by_project']}",
        f"Weekly commit breakdown: {period_data['weekly_commits']}",
        f"Papers published: {period_data['papers']}",
        f"Drafts by status: {period_data['drafts_by_status']}",
        f"Sync runs: {period_data['syncs']}",
    ])

    if goals:
        user_parts.append(f"\nGoals status: {goals}")
    if additional_instructions:
        user_parts.append(f"\nAdditional instructions: {additional_instructions}")

    user_prompt = "\n".join(user_parts)

    client = get_cloud_client()
    async with EgressRecorder(
        surface="worklog",
        service="worklog_service.generate_report",
        provider=type(client).__name__.lower().replace("client", ""),
        model=getattr(client, "_model", None) or getattr(client, "chat_model", None),
        user_id=None,
        project_id=None,
    ) as rec:
        rec.field("system_prompt", system_prompt)
        rec.field("metrics", str(period_data.get("commits_by_project", "")) + str(period_data.get("papers", "")) + str(period_data.get("drafts_by_status", "")) + str(period_data.get("syncs", "")))
        rec.field("drafts", str(period_data.get("drafts_by_status", "")))
        rec.field("papers", str(period_data.get("papers", "")))
        rec.field("goals", str(goals or ""))
        content = await client.complete(system_prompt, user_prompt)
    try:
        from app.services.event_stream import emit
        emit(
            "success",
            "worklog",
            f"{period_type} report generated",
            meta={"period": period_type},
        )
    except Exception:
        logger.exception("event emit failed (non-fatal)")
    return content


# ---------------------------------------------------------------------------
# DOCX export
# ---------------------------------------------------------------------------

def _generate_chart_png(weekly_data: List[Dict[str, Any]]) -> Optional[bytes]:
    """Render a bar chart of weekly commits as a PNG byte string."""
    if not weekly_data:
        return None

    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt

    weeks = [d["week"] for d in weekly_data]
    counts = [d["count"] for d in weekly_data]

    fig, ax = plt.subplots(figsize=(6, 3))
    ax.bar(range(len(weeks)), counts, color="#4F46E5")
    ax.set_xticks(range(len(weeks)))
    ax.set_xticklabels(weeks, rotation=45, ha="right", fontsize=7)
    ax.set_ylabel("Commits")
    ax.set_title("Weekly Commits")
    fig.tight_layout()

    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150)
    plt.close(fig)
    buf.seek(0)
    return buf.read()


def _add_formatted_text(para: Any, text: str) -> None:
    """Add text to a paragraph, handling **bold** markers."""
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = para.add_run(part[2:-2])
            run.bold = True
        else:
            para.add_run(part)


def export_to_docx(
    content: str,
    title: str,
    period_data: Dict[str, Any],
) -> bytes:
    """Convert markdown report to DOCX with formatting and an embedded chart."""
    from docx import Document
    from docx.shared import Inches, Pt

    doc = Document()

    # Title
    doc.add_heading(title, level=1)

    # Embed weekly commits chart if data exists
    chart_png = _generate_chart_png(period_data.get("weekly_commits", []))
    if chart_png:
        chart_stream = io.BytesIO(chart_png)
        doc.add_picture(chart_stream, width=Inches(5.5))
        doc.add_paragraph()  # spacer

    # Parse markdown content line by line
    lines = content.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]

        # Headings
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
            i += 1
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
            i += 1
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=2)
            i += 1
            continue

        # Markdown pipe table
        if "|" in line and line.strip().startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                stripped = lines[i].strip()
                # Skip separator rows like |---|---|
                if re.match(r"^\|[\s\-:|]+\|$", stripped):
                    i += 1
                    continue
                cells = [c.strip() for c in stripped.split("|")[1:-1]]
                table_lines.append(cells)
                i += 1

            if table_lines:
                num_cols = max(len(row) for row in table_lines)
                table = doc.add_table(rows=len(table_lines), cols=num_cols)
                table.style = "Table Grid"
                for row_idx, row_cells in enumerate(table_lines):
                    for col_idx, cell_text in enumerate(row_cells):
                        if col_idx < num_cols:
                            cell = table.rows[row_idx].cells[col_idx]
                            cell.text = cell_text
                            # Bold header row
                            if row_idx == 0:
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs:
                                        run.bold = True
            continue

        # Bullet points
        if line.strip().startswith("- "):
            para = doc.add_paragraph(style="List Bullet")
            _add_formatted_text(para, line.strip()[2:])
            i += 1
            continue

        # Regular paragraph (skip empty lines)
        text = line.strip()
        if text:
            para = doc.add_paragraph()
            _add_formatted_text(para, text)

        i += 1

    # Save to bytes
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
